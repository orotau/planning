from __future__ import print_function

from apiclient import discovery
from httplib2 import Http
from oauth2client import file, client, tools
import pprint
import collections
import more_itertools
from datetime import datetime
import iso8601 #http://pyiso8601.readthedocs.io/en/latest/

MATHS = "Maths"
TOPIC = "Topic"

doc_title = collections.namedtuple("doc_title", "year term week subject")
main_heading_data = collections.namedtuple("main_heading_data", "dayte, event_summary")


SCOPES = (
    'https://www.googleapis.com/auth/drive',
    'https://www.googleapis.com/auth/calendar.readonly',
)

# storage.json gets created when you run authorise.py
store = file.Storage('storage.json')
creds = store.get()
if not creds or creds.invalid:
    flow = client.flow_from_clientsecrets('client_secret.json', SCOPES)
    creds = tools.run_flow(flow, store)
DRIVE = discovery.build('drive', 'v3', http=creds.authorize(Http()))
CALENDAR = discovery.build('calendar', 'v3', http=creds.authorize(Http()))

# to run remove the X from the beginning of the folder ids
TOPIC_FOLDER_ID = "X1_IHOrfJsm9bNOrEkzIKsXQcBf5utVUfp" # term 3 2018
MATHS_FOLDER_ID = "X1_5iNW1G7_Sjfxhesx-LBAH-av8MsauO-" # term 3 2018


def get_calendar_id(term, year):
    calendars = []
    page_token = None
    while True:
        google_calendars = CALENDAR.calendarList().list(pageToken=page_token).execute()
        # returns a dict with four keys, very unlikely that we will need to go back again to the server
        # because we won't have that many calendars, but just in case.

        calendars.extend(google_calendars['items']) # create a list of dicts, 1 dict for each calendar

        page_token = google_calendars.get('nextPageToken')
        if not page_token:
            break

    for calendar in calendars:
        calendar_title = calendar['summary']
        # the assumption is that the 6th character will be the term .... 'Term X'
        # and the last 4 characters will be the year.
        if calendar_title[5] == str(term) and calendar_title.endswith(str(year)):
            return calendar['id']

    # no calendar has been found
    return None

def get_events_for_calendar(calendar_id):
    events = []
    page_token = None
    while True:
        google_events = CALENDAR.events().list(calendarId=calendar_id, pageToken=page_token).execute()

        events.extend(google_events['items'])
        page_token = google_events.get('nextPageToken')
        if not page_token:
            break

    return events

def get_documents_data(term, year, subject, calendar_events):
    '''
    We will return a dictionary with key = the named tuple ("doc_title", "year term week subject")
    and data = a list of the the named tuples ("main_heading_data", "dayte, event_summary")
    if there is no lesson on a dayte then both period and lesson_number will be set to 0
    '''
    documents_data = {}
    subject_keys = []

    all_day_events = [x for x in calendar_events if 'date' in x['start'] and 'date' in x['end'] and x['start']['date'] == x['end']['date']]
    # sort them in ascending order (to allow the week to be found)
    all_day_events = sorted(all_day_events, key=lambda k: k['start']['date'])

    line4_events = [x for x in calendar_events if x['summary'] and x['summary'].endswith("Line 4")]
    line5_events = [x for x in calendar_events if x['summary'] and x['summary'].endswith("Line 5")]

    # we are assuming Maths is Line 4 and Topic is Line 5
    if subject == MATHS:
        events_to_use = line4_events
    if subject == TOPIC:
        events_to_use = line5_events

    #We are assuming that the term starts on a Monday here (Not ok for term 1)
    assert datetime.strptime(all_day_events[0]['start']['date'], '%Y-%m-%d').weekday() == 0 #Monday

    for counter, weeks_all_day_events in enumerate(more_itertools.chunked(all_day_events, 5)):
        # each week
        week = counter + 1
        subject_keys.append(doc_title(year, term, week, subject))
        periods_for_week = []
        for all_day_event in weeks_all_day_events:
            # each day of the week
            for event in events_to_use:
                event_date = iso8601.parse_date(event['start']['dateTime']).date()
                all_day_event_date = datetime.strptime(all_day_event['start']['date'], '%Y-%m-%d').date()
                if event_date == all_day_event_date:                    
                    periods_for_week.append(main_heading_data(event_date, event['summary']))
        documents_data[subject_keys[counter]] = periods_for_week
    return(documents_data)

def create_documents(documents_data):
    # for each document, create the .docx, upload it as google doc
    for k,v in documents_data.items():
        from docx import Document

        new_document = Document()
        heading = str(k.year) + " - " + "Term " + str(k.term) + " - " + "Week " + str(k.week) + " - " + k.subject

        new_document.add_heading(heading, 0)
        new_document.add_heading("Summary", 1)
        new_document.add_paragraph('')
        for counter, heading_data in enumerate(v):
            date_to_use = datetime.strftime(heading_data.dayte, '%A %d %B')
            period_text = heading_data.event_summary[:8] # Period X
            lesson_text = "L" + str(counter + 1)
            heading = date_to_use + " - " + period_text + " - " + lesson_text
            new_document.add_heading(heading, 1)
            if k.subject == TOPIC:
                new_document.add_paragraph('')
            if k.subject == MATHS:
                # add a table for kākāriki and kiwikiwi
                table = new_document.add_table(rows=2, cols=2)
                table.style = 'TableGrid' #deprecated could break (it provides borders)
                # https://stackoverflow.com/questions/23725352/bold-a-table-cells
                table.cell(0, 0).paragraphs[0].add_run('kākāriki (green)').bold = True
                table.cell(0, 1).paragraphs[0].add_run('kiwikiwi (grey)').bold = True

        # to allow sorting properly use leading zero in week
        file_name = str(k.year) + " - " + "Term " + str(k.term) + " - " + "Week " + '{:02}'.format(k.week) + " - " + k.subject
        new_document.save(file_name + ".docx")

        #upload the document to google drive under the header folder
        if k.subject == MATHS:
            folder_id_to_use = MATHS_FOLDER_ID
        if k.subject == TOPIC:
            folder_id_to_use = TOPIC_FOLDER_ID

        # http://wescpy.blogspot.com/search/label/python3
        mimeType = "application/vnd.google-apps.document"
        body = {'name': file_name, 'mimeType': mimeType, 'parents': [folder_id_to_use]}
        DRIVE.files().create(body=body, media_body=file_name+'.docx', supportsTeamDrives=True, fields='id').execute().get('id')

def create_planning_skeletons(term, year):
    calendar_id = get_calendar_id(term, year)
    calendar_events = get_events_for_calendar(calendar_id)
    documents_data_maths = get_documents_data(term, year, MATHS, calendar_events)
    documents_data_topic = get_documents_data(term, year, TOPIC, calendar_events)
    create_documents(documents_data_maths)
    create_documents(documents_data_topic)

    

if __name__ == '__main__':

    import sys
    import argparse
    import ast
    import pprint

    # create the top-level parser
    parser = argparse.ArgumentParser()
    subparsers = parser.add_subparsers()

    # create the parser for get_calendar_id
    get_calendar_id_parser = subparsers.add_parser('get_calendar_id')
    get_calendar_id_parser.add_argument('term', type=int, choices = [1, 2, 3, 4])
    get_calendar_id_parser.add_argument('year', type=int)
    get_calendar_id_parser.set_defaults(function = get_calendar_id)

    # create the parser for get_events_for_calendar
    get_events_for_calendar_parser = subparsers.add_parser('get_events_for_calendar')
    get_events_for_calendar_parser.add_argument('id')
    get_events_for_calendar_parser.set_defaults(function = get_events_for_calendar)

    # create the parser for create_planning_skeletons
    create_planning_skeletons_parser = subparsers.add_parser('create_planning_skeletons')
    create_planning_skeletons_parser.add_argument('term', type=int, choices = [1, 2, 3, 4])
    create_planning_skeletons_parser.add_argument('year', type=int)
    create_planning_skeletons_parser.set_defaults(function = create_planning_skeletons)

    # parse the arguments
    arguments = parser.parse_args()
    arguments = vars(arguments) #convert from Namespace to dict

    #attempt to extract and then remove the function entry
    try:
        function_to_call = arguments['function'] 
    except KeyError:
        print ("You need a function name. Please type -h to get help")
        sys.exit()
    else:
        #remove the function entry as we are only passing arguments
        del arguments['function']
    
    if arguments:
        #remove any entries that have a value of 'None'
        #We are *assuming* that these are optional
        #We are doing this because we want the function definition to define
        #the defaults (NOT the function call)
        arguments = { k : v for k,v in arguments.items() if v is not None }

        #alter any string 'True' or 'False' to bools
        arguments = { k : ast.literal_eval(v) if v in ['True','False'] else v 
                                              for k,v in arguments.items() }       

    result = function_to_call(**arguments) #note **arguments works fine for empty dict {}
   
    pprint.pprint(result)
    print(len(result))

